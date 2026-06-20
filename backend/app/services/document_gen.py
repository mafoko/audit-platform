import io
from datetime import datetime
from typing import List, Any
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


class DocumentGenerator:
    def _add_heading(self, doc: Document, text: str, level: int = 1):
        heading = doc.add_heading(text, level=level)
        return heading

    def _add_table_row(self, table, *cells):
        row = table.add_row()
        for i, cell_text in enumerate(cells):
            row.cells[i].text = str(cell_text) if cell_text is not None else ""
        return row

    def generate_initiation_document(self, engagement: Any, controls: List[Any], evidence_requests: List[Any]) -> bytes:
        doc = Document()

        # Title
        title = doc.add_heading(f"Audit Initiation Document", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(f"Engagement: {engagement.title}")
        doc.add_paragraph(f"Client: {engagement.client_name}")
        doc.add_paragraph(f"Status: {engagement.status}")
        doc.add_paragraph(f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}")
        doc.add_paragraph("")

        # Scope & Objectives
        self._add_heading(doc, "1. Scope", 1)
        doc.add_paragraph(engagement.scope or "Not defined.")

        self._add_heading(doc, "2. Objectives", 1)
        doc.add_paragraph(engagement.objectives or "Not defined.")

        # Standards
        self._add_heading(doc, "3. Standards & Policies", 1)
        if engagement.standards:
            table = doc.add_table(rows=1, cols=3)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            hdr[0].text = "Name"
            hdr[1].text = "Type"
            hdr[2].text = "Content Preview"
            for std in engagement.standards:
                content_preview = (std.content or "")[:100] + ("..." if std.content and len(std.content) > 100 else "")
                self._add_table_row(table, std.name, std.type, content_preview)
        else:
            doc.add_paragraph("No standards attached.")

        # Controls
        self._add_heading(doc, "4. Controls", 1)
        if controls:
            table = doc.add_table(rows=1, cols=4)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            hdr[0].text = "Ref"
            hdr[1].text = "Title"
            hdr[2].text = "Category"
            hdr[3].text = "Description"
            for ctrl in controls:
                self._add_table_row(table, ctrl.control_ref, ctrl.title, ctrl.category or "", ctrl.description or "")
        else:
            doc.add_paragraph("No controls defined.")

        # Evidence Requests (PBC List)
        self._add_heading(doc, "5. Evidence Requests (PBC List)", 1)
        if evidence_requests:
            table = doc.add_table(rows=1, cols=5)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            hdr[0].text = "Ref"
            hdr[1].text = "Title"
            hdr[2].text = "Deadline (hrs)"
            hdr[3].text = "Due At"
            hdr[4].text = "Status"
            for er in evidence_requests:
                due_str = er.due_at.strftime("%Y-%m-%d %H:%M") if er.due_at else "TBD"
                self._add_table_row(table, er.request_ref, er.title, str(er.deadline_working_hours), due_str, er.status)
        else:
            doc.add_paragraph("No evidence requests defined.")

        # Timeline
        self._add_heading(doc, "6. Timeline", 1)
        start = engagement.start_date.strftime("%Y-%m-%d") if engagement.start_date else "TBD"
        end = engagement.end_date.strftime("%Y-%m-%d") if engagement.end_date else "TBD"
        doc.add_paragraph(f"Start Date: {start}")
        doc.add_paragraph(f"End Date: {end}")

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def generate_audit_report(self, report: Any, findings: List[Any], attestations: List[Any]) -> bytes:
        doc = Document()

        # Cover page
        title = doc.add_heading("AUDIT REPORT", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Report Title: {report.title}").alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Version: {report.version}").alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Status: {report.status}").alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}").alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_page_break()

        # Executive Summary
        self._add_heading(doc, "1. Executive Summary", 1)
        doc.add_paragraph(report.executive_summary or "No executive summary provided.")

        # Methodology
        self._add_heading(doc, "2. Methodology", 1)
        doc.add_paragraph(report.methodology or "No methodology documented.")

        # Scope
        self._add_heading(doc, "3. Scope", 1)
        doc.add_paragraph(report.scope_description or "No scope description provided.")

        # Findings Summary
        self._add_heading(doc, "4. Findings Summary", 1)
        rating_counts = {}
        for f in findings:
            rating_counts[f.risk_rating] = rating_counts.get(f.risk_rating, 0) + 1

        summary_table = doc.add_table(rows=1, cols=2)
        summary_table.style = "Table Grid"
        summary_table.rows[0].cells[0].text = "Risk Rating"
        summary_table.rows[0].cells[1].text = "Count"
        for rating in ["CRITICAL", "HIGH", "MEDIUM", "LOW", "INFORMATIONAL"]:
            row = summary_table.add_row()
            row.cells[0].text = rating
            row.cells[1].text = str(rating_counts.get(rating, 0))

        doc.add_paragraph("")

        # Detailed Findings
        self._add_heading(doc, "5. Detailed Findings", 1)
        if findings:
            for f in findings:
                self._add_heading(doc, f"{f.finding_ref}: {f.title}", 2)
                table = doc.add_table(rows=6, cols=2)
                table.style = "Table Grid"
                rows_data = [
                    ("Risk Rating", f.risk_rating),
                    ("Status", f.status),
                    ("Description", f.description or ""),
                    ("Root Cause", f.root_cause or ""),
                    ("Recommendation", f.recommendation or ""),
                    ("Management Response", f.management_response or "Awaiting response"),
                ]
                for i, (label, value) in enumerate(rows_data):
                    table.rows[i].cells[0].text = label
                    table.rows[i].cells[1].text = str(value)
                doc.add_paragraph("")
        else:
            doc.add_paragraph("No findings recorded.")

        # Sign-off / Attestation Page
        doc.add_page_break()
        self._add_heading(doc, "6. Sign-Off and Attestation", 1)
        if attestations:
            table = doc.add_table(rows=1, cols=4)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            hdr[0].text = "Name"
            hdr[1].text = "Role"
            hdr[2].text = "Date"
            hdr[3].text = "Note"
            for att in attestations:
                name = att.attested_by.full_name if hasattr(att, "attested_by") and att.attested_by else str(att.attested_by_id)
                self._add_table_row(table, name, att.role, att.attested_at.strftime("%Y-%m-%d"), att.signature_note or "")
        else:
            doc.add_paragraph("No attestations recorded.")

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()


doc_generator = DocumentGenerator()
